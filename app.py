import streamlit as st
import openpyxl
import json
import os
import io
import re
import shutil
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm, mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# ─── 설정 ───
st.set_page_config(page_title="품목제조보고서 자동생성", page_icon="📋", layout="wide")

# 한글 폰트 등록 (PDF용)
FONT_PATHS = [
    "C:/Windows/Fonts/malgun.ttf",
    "C:/Windows/Fonts/malgunbd.ttf",
]
FONT_REGISTERED = False
def register_fonts():
    global FONT_REGISTERED
    if FONT_REGISTERED:
        return
    try:
        pdfmetrics.registerFont(TTFont('Malgun', FONT_PATHS[0]))
        pdfmetrics.registerFont(TTFont('MalgunBold', FONT_PATHS[1]))
        FONT_REGISTERED = True
    except Exception as e:
        st.warning(f"폰트 등록 실패: {e}")

# ─── 별칭 매핑 (사용자가 확인한 매칭 저장) ───
ALIAS_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "alias_map.json")

def load_alias_map():
    if os.path.exists(ALIAS_PATH):
        with open(ALIAS_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def save_alias_map(alias_map):
    with open(ALIAS_PATH, "w", encoding="utf-8") as f:
        json.dump(alias_map, f, ensure_ascii=False, indent=2)

# ─── 원재료 DB 로드 ───
@st.cache_data
def load_ingredient_db():
    db_path = os.path.join(os.path.dirname(__file__), "ingredient_db.json")
    with open(db_path, "r", encoding="utf-8") as f:
        return json.load(f)

def normalize_name(s):
    """이름 정규화 (엑/액, 띠꼬/띠코 등 유사 표기 통일)"""
    s = s.strip()
    # 흔한 표기 차이 통일
    replacements = [
        ('액', '엑'), ('엑', '엑'),
        ('띠꼬', '띠코'), ('띠코', '띠코'),
        ('인스탄트', '인스턴트'), ('인스턴트', '인스턴트'),
        ('플리', '프리'), ('프리', '프리'),
    ]
    result = s
    for old, new in replacements:
        result = result.replace(old, new)
    return result

def find_ingredient(name, db):
    """원재료명으로 DB에서 검색 (유연한 매칭)"""
    name = name.strip()
    # 정확한 매칭
    if name in db:
        return db[name]
    # 띄어쓰기 제거 매칭
    no_space = name.replace(' ', '')
    for key in db:
        if key.replace(' ', '') == no_space:
            return db[key]
    # 정규화 매칭
    norm = normalize_name(name)
    for key in db:
        if normalize_name(key) == norm:
            return db[key]
    # 띄어쓰기 제거 + 정규화
    norm_no_space = normalize_name(no_space)
    for key in db:
        if normalize_name(key.replace(' ', '')) == norm_no_space:
            return db[key]
    # 부분 매칭
    for key in db:
        if key in name or name in key:
            return db[key]
    # 정규화된 부분 매칭
    for key in db:
        nk = normalize_name(key)
        if nk in norm or norm in nk:
            return db[key]
    # 괄호 제거 후 매칭
    clean = re.sub(r'[(\(（].*?[)\)）]', '', name).strip()
    if clean in db:
        return db[clean]
    for key in db:
        clean_key = re.sub(r'[(\(（].*?[)\)）]', '', key).strip()
        if clean_key == clean:
            return db[key]
    # 정규화 + 괄호 제거
    clean_norm = normalize_name(clean)
    for key in db:
        ck = normalize_name(re.sub(r'[(\(（].*?[)\)）]', '', key).strip())
        if ck == clean_norm:
            return db[key]
    # 별칭 매핑 확인
    alias_map = load_alias_map()
    if name in alias_map:
        mapped_key = alias_map[name]
        if mapped_key in db:
            return db[mapped_key]
    return None

def char_similarity(a, b):
    """두 문자열의 글자 단위 유사도 (0~1)"""
    if not a or not b:
        return 0.0
    a_set = set(a)
    b_set = set(b)
    intersection = a_set & b_set
    union = a_set | b_set
    jaccard = len(intersection) / len(union) if union else 0
    # 공통 글자 순서 보너스
    common_seq = 0
    j = 0
    for c in a:
        while j < len(b):
            if b[j] == c:
                common_seq += 1
                j += 1
                break
            j += 1
    seq_ratio = common_seq / max(len(a), len(b)) if max(len(a), len(b)) > 0 else 0
    return (jaccard + seq_ratio) / 2

def suggest_matches(name, db, top_n=5):
    """미매칭 원재료에 대해 DB에서 유사한 후보 추천"""
    name = name.strip()
    norm = normalize_name(name)
    clean = re.sub(r'[(\(（].*?[)\)）]', '', name).strip()
    clean_norm = normalize_name(clean)

    scores = []
    for key in db:
        key_clean = re.sub(r'[(\(（].*?[)\)）]', '', key).strip()
        key_norm = normalize_name(key_clean)

        # 여러 유사도 조합
        s1 = char_similarity(norm, normalize_name(key))
        s2 = char_similarity(clean_norm, key_norm)
        # 부분 포함 보너스
        bonus = 0
        if clean_norm in key_norm or key_norm in clean_norm:
            bonus = 0.3
        if len(clean_norm) >= 2 and clean_norm[:2] == key_norm[:2]:
            bonus += 0.1

        score = max(s1, s2) + bonus
        scores.append((key, score, db[key].get('성분', '')))

    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_n]

# ─── 엑셀 파싱 ───
def parse_recipe_excel(file):
    """원가 엑셀 파일에서 레시피 정보 추출"""
    wb = openpyxl.load_workbook(file, data_only=True)

    results = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        if ws.max_row is None or ws.max_row < 3:
            continue

        # 헤더 행 찾기 (원재료, 1배합, 백분율 등)
        header_row = None
        for row_idx in range(1, min(10, ws.max_row + 1)):
            for col_idx in range(1, min(10, ws.max_column + 1)):
                val = ws.cell(row=row_idx, column=col_idx).value
                if val and '원재료' in str(val):
                    header_row = row_idx
                    break
            if header_row:
                break

        if not header_row:
            continue

        # 제품명 찾기 (헤더 위)
        product_name = sheet_name
        for row_idx in range(1, header_row):
            for col_idx in range(1, min(10, ws.max_column + 1)):
                val = ws.cell(row=row_idx, column=col_idx).value
                if val and isinstance(val, str) and len(val) > 1 and '단위' not in val:
                    product_name = val.strip()
                    break

        # 컬럼 인덱스 찾기
        col_map = {}
        for col_idx in range(1, min(10, ws.max_column + 1)):
            val = ws.cell(row=header_row, column=col_idx).value
            if val:
                val = str(val).strip()
                if '원재료' in val:
                    col_map['name'] = col_idx
                elif '1배합' in val or val == '배합':
                    col_map['amount'] = col_idx
                elif '백분율' in val:
                    col_map['percent'] = col_idx

        if 'name' not in col_map:
            continue

        ingredients = []
        for row_idx in range(header_row + 1, ws.max_row + 1):
            name_val = ws.cell(row=row_idx, column=col_map['name']).value
            if not name_val or not isinstance(name_val, str):
                continue
            name_val = name_val.strip()
            if name_val in ('합계', '합 계', '계', '소계'):
                break
            if not name_val:
                continue

            amount = ws.cell(row=row_idx, column=col_map.get('amount', 2)).value or 0
            percent = ws.cell(row=row_idx, column=col_map.get('percent', 7)).value or 0

            try:
                amount = float(amount)
                percent = float(percent)
            except (ValueError, TypeError):
                continue

            ingredients.append({
                'name': name_val,
                'amount': amount,
                'percent': round(percent, 2)
            })

        if ingredients:
            results.append({
                'product_name': product_name,
                'sheet_name': sheet_name,
                'ingredients': ingredients,
                'total_amount': sum(i['amount'] for i in ingredients)
            })

    return results

# ─── Word 문서 생성 ───
def set_cell_text(cell, text, bold=False, font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """셀에 텍스트 설정"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(str(text))
    run.font.size = Pt(font_size)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.bold = bold

def set_cell_shading(cell, color="D9E2F3"):
    """셀 배경색 설정"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_table_border(table):
    """테이블 전체 테두리 설정"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

def generate_doc1_manufacturing(recipe, options, db):
    """1. 제조방법설명서 생성"""
    doc = Document()

    # 여백 설정
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 제목
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('제 조 방 법 설 명 서')
    run.font.size = Pt(18)
    run.bold = True
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    doc.add_paragraph()

    # 테이블 생성
    ingredients = recipe['ingredients']
    product_name = recipe['product_name']

    # 원재료명, 성분명 및 배합비율 문자열 생성 (표시명 = 식품유형명 사용)
    ingredient_parts = []
    for ing in sorted(ingredients, key=lambda x: x['percent'], reverse=True):
        info = find_ingredient(ing['name'], db)
        if info:
            display_name = info.get('표시명', info.get('원재료명', ing['name']))
        else:
            display_name = ing['name']

        ingredient_parts.append(f"{display_name} {ing['percent']}%")

    ingredient_str = ", ".join(ingredient_parts)

    # 제조방법 생성
    ing_names = [i['name'] for i in ingredients if i['name'] != '정제수']
    method_names = ",".join(ing_names[:5])
    if len(ing_names) > 5:
        method_names += " 등"

    storage_method = options.get('storage', '냉동')
    bake_temp = options.get('bake_temp', '180')
    bake_time = options.get('bake_time', '30~40')
    ferment1_time = options.get('ferment1_time', '40~50')
    ferment2_time = options.get('ferment2_time', '50~60')
    divide_weight = options.get('divide_weight', '320')

    if storage_method == '냉동':
        method_text = (
            f"{method_names},정제수를 계량후 믹싱을 한다 "
            f"(1)을 1차발효 {ferment1_time}분을 한다 "
            f"(2)를 {divide_weight}g 분할후 둥글리기를 한다 "
            f"(3)을 팬에 넣고 2차발효 {ferment2_time}분을 한다 "
            f"(4)를 {bake_temp}℃로 예열된 오븐에 넣고 {bake_time}분간 가열을 한다 "
            f"(5)를 탈팬후 식혀 포장 냉동을 한다"
        )
    elif storage_method == '냉장':
        method_text = (
            f"{method_names},정제수를 계량후 믹싱을 한다 "
            f"(1)을 1차발효 {ferment1_time}분을 한다 "
            f"(2)를 {divide_weight}g 분할후 둥글리기를 한다 "
            f"(3)을 팬에 넣고 2차발효 {ferment2_time}분을 한다 "
            f"(4)를 {bake_temp}℃로 예열된 오븐에 넣고 {bake_time}분간 가열을 한다 "
            f"(5)를 탈팬후 식혀 포장 냉장을 한다"
        )
    else:
        method_text = (
            f"{method_names},정제수를 계량후 믹싱을 한다 "
            f"(1)을 1차발효 {ferment1_time}분을 한다 "
            f"(2)를 {divide_weight}g 분할후 둥글리기를 한다 "
            f"(3)을 팬에 넣고 2차발효 {ferment2_time}분을 한다 "
            f"(4)를 {bake_temp}℃로 예열된 오븐에 넣고 {bake_time}분간 가열을 한다 "
            f"(5)를 탈팬후 식혀 포장을 한다"
        )

    storage_map = {
        '냉동': '-18℃이하 냉동보관',
        '냉장': '0~10℃ 냉장보관',
        '실온': '직사광선을 피하고 서늘한 곳에 보관'
    }

    usage_map = {
        '냉동': '소비자 판매용 / 해동 후 섭취',
        '냉장': '소비자 판매용 / 냉장 보관 후 섭취',
        '실온': '소비자 판매용 / 개봉 후 섭취'
    }

    shelf_life = options.get('shelf_life', '9개월')

    rows_data = [
        ('식품군', '과자류, 빵류 또는 떡류'),
        ('식품의 유형', '빵류'),
        ('제품명', product_name),
        ('원재료명, 성분명 및\n배합비율(%)', ingredient_str),
        ('제조방법', method_text),
        ('용도용법', usage_map.get(storage_method, usage_map['냉동'])),
        ('보관방법', storage_map.get(storage_method, storage_map['냉동'])),
        ('포장재질', options.get('packaging_material', 'PP, PS, PE, PET')),
        ('포장방법 및 포장단위', f"밀봉 / {options.get('packaging_unit', '10g~5kg')}"),
        ('성상', '고유의 향미를 가지고 있으며 이미, 이취가 없음'),
        ('소비기한', f"제조일로부터 {shelf_life}까지"),
        ('기타', ''),
    ]

    table = doc.add_table(rows=len(rows_data) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table)

    # 헤더
    hdr = table.rows[0]
    set_cell_text(hdr.cells[0], '구분', bold=True, font_size=11)
    set_cell_text(hdr.cells[1], '신고내역', bold=True, font_size=11)
    set_cell_shading(hdr.cells[0])
    set_cell_shading(hdr.cells[1])

    # 컬럼 너비
    for row in table.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(12)

    for i, (label, value) in enumerate(rows_data):
        row = table.rows[i + 1]
        set_cell_text(row.cells[0], label, bold=True, font_size=10)
        set_cell_text(row.cells[1], value, font_size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_shading(row.cells[0], "F2F2F2")

    return doc

def generate_doc2_ingredients(recipe, options, db):
    """2. 원료성분 및 배합비율 생성"""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # 제목
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('원료성분 및 배합비율')
    run.font.size = Pt(18)
    run.bold = True
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    doc.add_paragraph()

    ingredients = recipe['ingredients']
    # 백분율 내림차순 정렬
    ingredients_sorted = sorted(ingredients, key=lambda x: x['percent'], reverse=True)

    # 헤더 + 재료 + 합계 (빈 행 없이 딱 맞게)
    table = doc.add_table(rows=1 + len(ingredients_sorted) + 1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table)

    # 헤더
    headers = ['원재료명', '배합비율(%)', '성                 분', '비고']
    for j, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], h, bold=True, font_size=11)
        set_cell_shading(table.rows[0].cells[j])

    # 컬럼 너비
    widths = [Cm(3.5), Cm(2.5), Cm(8.5), Cm(2.5)]
    for row in table.rows:
        for j, w in enumerate(widths):
            row.cells[j].width = w

    # 데이터 행
    for i, ing in enumerate(ingredients_sorted):
        row = table.rows[i + 1]
        info = find_ingredient(ing['name'], db)

        if info:
            display_name = info.get('원재료명', ing['name'])
            composition = info.get('성분', '')
            origin = info.get('원산지', '')
        else:
            display_name = ing['name']
            composition = ''
            origin = ''

        set_cell_text(row.cells[0], display_name, font_size=10)
        set_cell_text(row.cells[1], f"{ing['percent']:.2f}", font_size=10)
        set_cell_text(row.cells[2], composition, font_size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row.cells[3], origin, font_size=9)

    # 합계 행
    sum_row = table.rows[-1]
    set_cell_text(sum_row.cells[0], '합 계', bold=True, font_size=11)
    set_cell_text(sum_row.cells[1], '100', bold=True, font_size=11)
    set_cell_text(sum_row.cells[2], '', font_size=10)
    set_cell_text(sum_row.cells[3], '', font_size=10)
    set_cell_shading(sum_row.cells[0], "F2F2F2")
    set_cell_shading(sum_row.cells[1], "F2F2F2")

    return doc

def generate_doc3_shelf_life(recipe, options, db):
    """3. 소비기한설정사유서 생성"""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    product_name = recipe['product_name']
    storage_method = options.get('storage', '냉동')
    shelf_life = options.get('shelf_life', '9개월')
    packaging_material = options.get('packaging_material', 'PP, PS, PE, PET')

    # 주성분 (상위 3개)
    top3 = sorted(recipe['ingredients'], key=lambda x: x['percent'], reverse=True)[:3]
    top3_names = []
    for ing in top3:
        info = find_ingredient(ing['name'], db)
        if info:
            top3_names.append(info.get('원재료명', ing['name']))
        else:
            top3_names.append(ing['name'])
    main_ingredients = ", ".join(top3_names)

    storage_check = {'실온': '', '상온': '', '냉장': '', '냉동': '', '기타': ''}
    storage_check[storage_method] = 'O'

    # 비교제품 정보
    compare_product = options.get('compare_product', '통밀식빵(냉동)')
    compare_company = options.get('compare_company', '더브레드블루')

    # 제목
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('소비기한 설정 사유서')
    run.font.size = Pt(18)
    run.bold = True
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    doc.add_paragraph()

    # 기본 정보 테이블
    table1 = doc.add_table(rows=5, cols=2)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table1)

    info_rows = [
        ('제   품   명', product_name),
        ('식품·축산물의 유형', '빵류'),
        ('보존 및 유통 방법',
         f"실온( {storage_check['실온']} ) / 상온( {storage_check['상온']} ) / "
         f"냉장( {storage_check['냉장']} ) / 냉동( {storage_check['냉동']} ) / 기타( {storage_check['기타']} )"),
        ('소   비   기   한', f"제조일로부터 {shelf_life}까지"),
        ('실험수행기관종류', '자사(   ) / 의뢰(   ) / 생략( O )'),
    ]

    for row in table1.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(12)

    for i, (label, value) in enumerate(info_rows):
        set_cell_text(table1.rows[i].cells[0], label, bold=True, font_size=10)
        set_cell_text(table1.rows[i].cells[1], value, font_size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_shading(table1.rows[i].cells[0], "F2F2F2")

    doc.add_paragraph()

    # 소비기한 설정근거 제목
    p = doc.add_paragraph()
    run = p.add_run('소비기한 설정근거')
    run.font.size = Pt(12)
    run.bold = True
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # 설정근거 본문
    reason_text = (
        f"제품의 원료 및 보존 특성\n"
        f"본 제품의 주성분은 {main_ingredients}이며, "
        f"포장재질이 {packaging_material}으로 밀봉포장되어 "
        f"외부의 공기 및 습기가 침투하지 못하므로 미생물의 생육이 억제됨."
    )
    p2 = doc.add_paragraph(reason_text)
    for run in p2.runs:
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'

    doc.add_paragraph()

    # 유사제품 비교 제목
    p3 = doc.add_paragraph()
    run = p3.add_run('유사제품 비교')
    run.font.size = Pt(12)
    run.bold = True
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # 비교 테이블
    compare_table = doc.add_table(rows=11, cols=3)
    compare_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(compare_table)

    for row in compare_table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(6.5)
        row.cells[2].width = Cm(6.5)

    compare_data = [
        ('구분', f'신규제품(현제품)', f'기존 유통제품(비교제품)'),
        ('제품명', product_name, compare_product),
        ('제조사', '더브레드블루', compare_company),
        ('식품유형', '빵류', '빵류'),
        ('성상', '고유의 향미를 가지고 있으며, 이미, 이취가 없음.', '고유의 향미를 가지고 있으며, 이미, 이취가 없음.'),
        ('포장재질', packaging_material, packaging_material),
        ('포장방법', '밀봉', '밀봉'),
        ('보존 및 유통온도', storage_method, storage_method),
        ('보존료 사용여부', '미사용', '미사용'),
        ('유탕·유처리', '미처리', '미처리'),
        ('유통기한', f"제조일로부터 {shelf_life}까지", f"제조일로부터 {shelf_life}까지"),
    ]

    for i, (col0, col1, col2) in enumerate(compare_data):
        row = compare_table.rows[i]
        set_cell_text(row.cells[0], col0, bold=(i == 0), font_size=10)
        set_cell_text(row.cells[1], col1, bold=(i == 0), font_size=10)
        set_cell_text(row.cells[2], col2, bold=(i == 0), font_size=10)
        if i == 0:
            for j in range(3):
                set_cell_shading(row.cells[j])
        else:
            set_cell_shading(row.cells[0], "F2F2F2")

    doc.add_paragraph()

    # 종합 판단
    p4 = doc.add_paragraph()
    run = p4.add_run('종합 판단')
    run.font.size = Pt(12)
    run.bold = True
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    conclusion = (
        f"{packaging_material} 밀봉포장으로 미생물의 생육이 억제되는 점과 "
        f"본 제품과 제품특성이 유사한 기존 유통제품의 소비기한이 "
        f"제조일로부터 {shelf_life}까지인 점을 감안하여 "
        f"'식품 등의 소비기한 설정기준'에 따라 소비기한 설정실험을 생략하고 "
        f"본 제품의 소비기한을 제조일로부터 {shelf_life}까지로 설정합니다."
    )
    p5 = doc.add_paragraph(conclusion)
    for run in p5.runs:
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'

    doc.add_paragraph()
    doc.add_paragraph()

    # 서명란
    today = datetime.now()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sig.add_run(
        f"상기와 같이 소비기한 설정 사유서를 제출합니다.\n\n"
        f"        {today.year}년    {today.month:02d}월    {today.day:02d}일\n\n"
        f"        제출인 :     문동진   (인)"
    )
    run.font.size = Pt(11)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    return doc

# ─── PDF 생성 ───
def generate_pdf(doc_type, recipe, options, db):
    """PDF 생성 (HWP 양식 재현)"""
    register_fonts()

    buffer = io.BytesIO()
    pdf = SimpleDocTemplate(
        buffer, pagesize=A4,
        topMargin=2.5*cm, bottomMargin=2.5*cm,
        leftMargin=2*cm, rightMargin=2*cm
    )

    styles = getSampleStyleSheet()

    # 한글 스타일
    style_title = ParagraphStyle(
        'KorTitle', parent=styles['Title'],
        fontName='Malgun', fontSize=16, alignment=1,
        spaceAfter=20
    )
    style_normal = ParagraphStyle(
        'KorNormal', parent=styles['Normal'],
        fontName='Malgun', fontSize=9, leading=14
    )
    style_bold = ParagraphStyle(
        'KorBold', parent=styles['Normal'],
        fontName='MalgunBold', fontSize=9, leading=14, alignment=1
    )
    style_center = ParagraphStyle(
        'KorCenter', parent=styles['Normal'],
        fontName='Malgun', fontSize=9, leading=14, alignment=1
    )
    style_small = ParagraphStyle(
        'KorSmall', parent=styles['Normal'],
        fontName='Malgun', fontSize=8, leading=11
    )

    elements = []
    product_name = recipe['product_name']
    ingredients = recipe['ingredients']
    storage_method = options.get('storage', '냉동')
    shelf_life = options.get('shelf_life', '9개월')
    packaging_material = options.get('packaging_material', 'PP, PS, PE, PET')

    if doc_type == 'manufacturing':
        # 1. 제조방법설명서 PDF
        elements.append(Paragraph('제 조 방 법 설 명 서', style_title))
        elements.append(Spacer(1, 20))

        # 원재료명 배합비율 문자열
        ing_parts = []
        for ing in sorted(ingredients, key=lambda x: x['percent'], reverse=True):
            info = find_ingredient(ing['name'], db)
            display = info.get('표시명', info.get('원재료명', ing['name'])) if info else ing['name']
            ing_parts.append(f"{display} {ing['percent']}%")
        ing_str = ", ".join(ing_parts)

        # 제조방법
        ing_names = [i['name'] for i in ingredients if i['name'] != '정제수']
        method_names = ",".join(ing_names[:5])
        if len(ing_names) > 5:
            method_names += " 등"

        bake_temp = options.get('bake_temp', '180')
        bake_time = options.get('bake_time', '30~40')
        ferment1 = options.get('ferment1_time', '40~50')
        ferment2 = options.get('ferment2_time', '50~60')
        div_wt = options.get('divide_weight', '320')

        storage_suffix = {'냉동': '포장 냉동을 한다', '냉장': '포장 냉장을 한다', '실온': '포장을 한다'}
        method_text = (
            f"{method_names},정제수를 계량후 믹싱을 한다 "
            f"(1)을 1차발효 {ferment1}분을 한다 "
            f"(2)를 {div_wt}g 분할후 둥글리기를 한다 "
            f"(3)을 팬에 넣고 2차발효 {ferment2}분을 한다 "
            f"(4)를 {bake_temp}℃로 예열된 오븐에 넣고 {bake_time}분간 가열을 한다 "
            f"(5)를 탈팬후 식혀 {storage_suffix.get(storage_method, storage_suffix['냉동'])}"
        )

        storage_desc = {'냉동': '-18℃이하 냉동보관', '냉장': '0~10℃ 냉장보관', '실온': '직사광선을 피하고 서늘한 곳에 보관'}
        usage_desc = {'냉동': '소비자 판매용 / 해동 후 섭취', '냉장': '소비자 판매용 / 냉장 보관 후 섭취', '실온': '소비자 판매용 / 개봉 후 섭취'}

        data = [
            [Paragraph('구분', style_bold), Paragraph('신고내역', style_bold)],
            [Paragraph('식품군', style_center), Paragraph('과자류, 빵류 또는 떡류', style_normal)],
            [Paragraph('식품의 유형', style_center), Paragraph('빵류', style_normal)],
            [Paragraph('제품명', style_center), Paragraph(product_name, style_normal)],
            [Paragraph('원재료명, 성분명<br/>및 배합비율(%)', style_center), Paragraph(ing_str, style_small)],
            [Paragraph('제조방법', style_center), Paragraph(method_text, style_small)],
            [Paragraph('용도용법', style_center), Paragraph(usage_desc.get(storage_method, ''), style_normal)],
            [Paragraph('보관방법', style_center), Paragraph(storage_desc.get(storage_method, ''), style_normal)],
            [Paragraph('포장재질', style_center), Paragraph(packaging_material, style_normal)],
            [Paragraph('포장방법 및<br/>포장단위', style_center), Paragraph(f"밀봉 / {options.get('packaging_unit', '10g~5kg')}", style_normal)],
            [Paragraph('성상', style_center), Paragraph('고유의 향미를 가지고 있으며 이미, 이취가 없음', style_normal)],
            [Paragraph('소비기한', style_center), Paragraph(f'제조일로부터 {shelf_life}까지', style_normal)],
            [Paragraph('기타', style_center), Paragraph('', style_normal)],
        ]

        t = Table(data, colWidths=[4.5*cm, 12.5*cm])
        t.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#D9E2F3')),
            ('BACKGROUND', (0, 1), (0, -1), colors.HexColor('#F2F2F2')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ALIGN', (0, 0), (0, -1), 'CENTER'),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('LEFTPADDING', (1, 1), (1, -1), 8),
        ]))
        elements.append(t)

    elif doc_type == 'ingredients':
        # 2. 원료성분 및 배합비율 PDF
        elements.append(Paragraph('원료성분 및 배합비율', style_title))
        elements.append(Spacer(1, 20))

        sorted_ings = sorted(ingredients, key=lambda x: x['percent'], reverse=True)

        data = [[
            Paragraph('원재료명', style_bold),
            Paragraph('배합비율(%)', style_bold),
            Paragraph('성                 분', style_bold),
            Paragraph('비고', style_bold),
        ]]

        for ing in sorted_ings:
            info = find_ingredient(ing['name'], db)
            display_name = info.get('원재료명', ing['name']) if info else ing['name']
            comp = info.get('성분', '') if info else ''
            origin = info.get('원산지', '') if info else ''
            data.append([
                Paragraph(display_name, style_center),
                Paragraph(f"{ing['percent']:.2f}", style_center),
                Paragraph(comp, style_small),
                Paragraph(origin, style_center),
            ])

        # 합계
        data.append([
            Paragraph('합 계', style_bold),
            Paragraph('100', style_bold),
            Paragraph('', style_center),
            Paragraph('', style_center),
        ])

        t = Table(data, colWidths=[3.5*cm, 2.5*cm, 8*cm, 3*cm])
        t.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#D9E2F3')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('TOPPADDING', (0, 0), (-1, -1), 3),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ]))
        elements.append(t)

    elif doc_type == 'shelf_life':
        # 3. 소비기한설정사유서 PDF
        elements.append(Paragraph('소비기한 설정 사유서', style_title))
        elements.append(Spacer(1, 15))

        storage_check = {'실온': '   ', '상온': '   ', '냉장': '   ', '냉동': '   ', '기타': '   '}
        storage_check[storage_method] = ' O '

        data1 = [
            [Paragraph('제   품   명', style_bold), Paragraph(product_name, style_normal)],
            [Paragraph('식품·축산물의 유형', style_bold), Paragraph('빵류', style_normal)],
            [Paragraph('보존 및 유통 방법', style_bold),
             Paragraph(f"실온({storage_check['실온']}) / 상온({storage_check['상온']}) / 냉장({storage_check['냉장']}) / 냉동({storage_check['냉동']}) / 기타({storage_check['기타']})", style_normal)],
            [Paragraph('소   비   기   한', style_bold), Paragraph(f'제조일로부터 {shelf_life}까지', style_normal)],
            [Paragraph('실험수행기관종류', style_bold), Paragraph('자사(   ) / 의뢰(   ) / 생략( O )', style_normal)],
        ]

        t1 = Table(data1, colWidths=[5*cm, 12*cm])
        t1.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#F2F2F2')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('LEFTPADDING', (1, 0), (1, -1), 8),
        ]))
        elements.append(t1)
        elements.append(Spacer(1, 15))

        # 설정근거
        elements.append(Paragraph('<b>소비기한 설정근거</b>', ParagraphStyle('h2', parent=style_normal, fontName='MalgunBold', fontSize=11, spaceAfter=10)))

        top3 = sorted(ingredients, key=lambda x: x['percent'], reverse=True)[:3]
        top3_names = []
        for ing in top3:
            info = find_ingredient(ing['name'], db)
            top3_names.append(info.get('원재료명', ing['name']) if info else ing['name'])
        main_ings = ", ".join(top3_names)

        elements.append(Paragraph(
            f"<b>제품의 원료 및 보존 특성</b><br/>"
            f"본 제품의 주성분은 {main_ings}이며, "
            f"포장재질이 {packaging_material}으로 밀봉포장되어 "
            f"외부의 공기 및 습기가 침투하지 못하므로 미생물의 생육이 억제됨.",
            style_normal
        ))
        elements.append(Spacer(1, 15))

        # 비교 테이블
        elements.append(Paragraph('<b>유사제품 비교</b>', ParagraphStyle('h2b', parent=style_normal, fontName='MalgunBold', fontSize=11, spaceAfter=10)))

        compare_product = options.get('compare_product', '통밀식빵(냉동)')

        cmp_data = [
            [Paragraph('구분', style_bold), Paragraph('신규제품(현제품)', style_bold), Paragraph('기존 유통제품(비교제품)', style_bold)],
            [Paragraph('제품명', style_center), Paragraph(product_name, style_center), Paragraph(compare_product, style_center)],
            [Paragraph('제조사', style_center), Paragraph('더브레드블루', style_center), Paragraph(options.get('compare_company', '더브레드블루'), style_center)],
            [Paragraph('식품유형', style_center), Paragraph('빵류', style_center), Paragraph('빵류', style_center)],
            [Paragraph('성상', style_center), Paragraph('고유의 향미를 가지고 있으며,<br/>이미, 이취가 없음.', style_small), Paragraph('고유의 향미를 가지고 있으며,<br/>이미, 이취가 없음.', style_small)],
            [Paragraph('포장재질', style_center), Paragraph(packaging_material, style_center), Paragraph(packaging_material, style_center)],
            [Paragraph('포장방법', style_center), Paragraph('밀봉', style_center), Paragraph('밀봉', style_center)],
            [Paragraph('보존 및 유통온도', style_center), Paragraph(storage_method, style_center), Paragraph(storage_method, style_center)],
            [Paragraph('보존료 사용여부', style_center), Paragraph('미사용', style_center), Paragraph('미사용', style_center)],
            [Paragraph('유탕·유처리', style_center), Paragraph('미처리', style_center), Paragraph('미처리', style_center)],
            [Paragraph('유통기한', style_center), Paragraph(f'제조일로부터 {shelf_life}까지', style_center), Paragraph(f'제조일로부터 {shelf_life}까지', style_center)],
        ]

        t2 = Table(cmp_data, colWidths=[4*cm, 6.5*cm, 6.5*cm])
        t2.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#D9E2F3')),
            ('BACKGROUND', (0, 1), (0, -1), colors.HexColor('#F2F2F2')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 0), (-1, -1), 3),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ]))
        elements.append(t2)
        elements.append(Spacer(1, 15))

        # 종합 판단
        elements.append(Paragraph('<b>종합 판단</b>', ParagraphStyle('h2c', parent=style_normal, fontName='MalgunBold', fontSize=11, spaceAfter=10)))
        elements.append(Paragraph(
            f"{packaging_material} 밀봉포장으로 미생물의 생육이 억제되는 점과 "
            f"본 제품과 제품특성이 유사한 기존 유통제품의 소비기한이 "
            f"제조일로부터 {shelf_life}까지인 점을 감안하여 "
            f"'식품 등의 소비기한 설정기준'에 따라 소비기한 설정실험을 생략하고 "
            f"본 제품의 소비기한을 제조일로부터 {shelf_life}까지로 설정합니다.",
            style_normal
        ))
        elements.append(Spacer(1, 30))

        today = datetime.now()
        elements.append(Paragraph(
            f"상기와 같이 소비기한 설정 사유서를 제출합니다.<br/><br/>"
            f"        {today.year}년    {today.month:02d}월    {today.day:02d}일<br/><br/>"
            f"        제출인 :     문동진   (인)",
            ParagraphStyle('sig', parent=style_normal, fontSize=11, alignment=1, leading=20)
        ))

    pdf.build(elements)
    buffer.seek(0)
    return buffer

# ─── HWP 생성 (한글 자동화) ───
HWP_TEMPLATES = {
    'doc1': r'c:/Users/Moon/OneDrive/03. 양지희_품질/2. 품목제조보고/1. 빵류/25년/호라산밀식빵/1. 제조방법설명서 - 호라산밀식빵.hwp',
    'doc2': r'c:/Users/Moon/OneDrive/03. 양지희_품질/2. 품목제조보고/1. 빵류/25년/호라산밀식빵/2. 원료성분 및 배합비율(호라산밀식빵).hwp',
    'doc3': r'c:/Users/Moon/OneDrive/03. 양지희_품질/2. 품목제조보고/1. 빵류/25년/호라산밀식빵/3. 소비기한설정사유서(호라산밀식빵).hwp',
}

def hwp_replace(hwp, find_str, replace_str):
    """HWP 문서에서 찾아바꾸기"""
    hwp.HAction.GetDefault('AllReplace', hwp.HParameterSet.HFindReplace.HSet)
    pset = hwp.HParameterSet.HFindReplace
    pset.FindString = str(find_str)
    pset.ReplaceString = str(replace_str)
    pset.IgnoreMessage = 1
    pset.Direction = 0
    pset.ReplaceMode = 1
    hwp.HAction.Execute('AllReplace', pset.HSet)

def hwp_clear_table_and_fill(hwp, ingredients_data):
    """원료성분 테이블: 기존 데이터 지우고 새 데이터로 채우기"""
    # 첫번째 표 찾기
    hwp.HAction.GetDefault('ShapeObjTableSelCell', hwp.HParameterSet.HShapeObject.HSet)

    # 커서를 문서 처음으로
    hwp.MovePos(2)  # 문서 처음

    # 표 안으로 이동 — 첫 번째 표 찾기
    hwp.HAction.Run('MoveNextParaBegin')
    hwp.HAction.Run('MoveNextParaBegin')
    hwp.HAction.Run('TableCellBlock')

    # 표의 각 셀을 순회하며 데이터 입력
    # 헤더 행(1행)은 건너뛰고, 2행부터 데이터 입력

    # 2행 1열로 이동 (첫 번째 데이터 행)
    hwp.HAction.Run('TableLowerCell')  # 아래 셀로

    for i, ing in enumerate(ingredients_data):
        # 원재료명
        hwp.HAction.Run('Select')  # 셀 전체 선택
        hwp.HAction.GetDefault('InsertText', hwp.HParameterSet.HInsertText.HSet)
        hwp.HParameterSet.HInsertText.Text = ing['name']
        hwp.HAction.Execute('InsertText', hwp.HParameterSet.HInsertText.HSet)

        # 배합비율
        hwp.HAction.Run('TableRightCell')
        hwp.HAction.GetDefault('InsertText', hwp.HParameterSet.HInsertText.HSet)
        hwp.HParameterSet.HInsertText.Text = f"{ing['percent']:.2f}"
        hwp.HAction.Execute('InsertText', hwp.HParameterSet.HInsertText.HSet)

        # 성분
        hwp.HAction.Run('TableRightCell')
        hwp.HAction.GetDefault('InsertText', hwp.HParameterSet.HInsertText.HSet)
        hwp.HParameterSet.HInsertText.Text = ing.get('composition', '')
        hwp.HAction.Execute('InsertText', hwp.HParameterSet.HInsertText.HSet)

        # 비고
        hwp.HAction.Run('TableRightCell')
        hwp.HAction.GetDefault('InsertText', hwp.HParameterSet.HInsertText.HSet)
        hwp.HParameterSet.HInsertText.Text = ing.get('origin', '')
        hwp.HAction.Execute('InsertText', hwp.HParameterSet.HInsertText.HSet)

        # 다음 행
        if i < len(ingredients_data) - 1:
            hwp.HAction.Run('TableLowerCell')
            hwp.HAction.Run('TableCellBlockCol')
            hwp.HAction.Run('TableLeftCell')

def generate_hwp_files(recipe, options, db):
    """HWP 파일 3종 생성 — 기존 템플릿 기반 찾아바꾸기"""
    import win32com.client
    import pythoncom

    pythoncom.CoInitialize()

    product_name = recipe['product_name']
    ingredients = recipe['ingredients']
    storage_method = options.get('storage', '냉동')
    shelf_life = options.get('shelf_life', '9개월')
    packaging_material = options.get('packaging_material', 'PP, PS, PE, PET')

    # 정렬된 원재료 (백분율 내림차순)
    sorted_ings = sorted(ingredients, key=lambda x: x['percent'], reverse=True)

    # 제조방법설명서용 원재료명 문자열
    ing_parts = []
    for ing in sorted_ings:
        info = find_ingredient(ing['name'], db)
        display = info.get('표시명', info.get('원재료명', ing['name'])) if info else ing['name']
        ing_parts.append(f"{display} {ing['percent']}%")
    ingredient_str = ", ".join(ing_parts)

    # 제조방법 문자열
    ing_names_list = []
    for i in sorted_ings:
        if i['name'] != '정제수':
            info = find_ingredient(i['name'], db)
            ing_names_list.append(info.get('원재료명', i['name']) if info else i['name'])
    method_names = ",".join(ing_names_list[:7])

    storage_suffix = {'냉동': '포장 냉동을 한다', '냉장': '포장 냉장을 한다', '실온': '포장을 한다'}
    bake_temp = options.get('bake_temp', '180')
    bake_time = options.get('bake_time', '30~40')
    ferment1 = options.get('ferment1_time', '40~50')
    ferment2 = options.get('ferment2_time', '50~60')
    div_wt = options.get('divide_weight', '320')

    method_text = (
        f"{method_names},정제수를 계량후 믹싱을 한다\r\n"
        f"(1)을 1차발효 {ferment1}분을 한다\r\n"
        f"(2)를 {div_wt}g 분할후 둥글리기를 한다\r\n"
        f"(3)을 팬에 넣고 2차발효 {ferment2}분을 한다\r\n"
        f"(4)를 {bake_temp}℃로 예열된 오븐에 넣고 {bake_time}분간 가열을 한다\r\n"
        f"(5)를 탈팬후 식혀 {storage_suffix.get(storage_method, storage_suffix['냉동'])}"
    )

    storage_map = {'냉동': '-18℃이하 냉동보관', '냉장': '0~10℃ 냉장보관', '실온': '직사광선을 피하고 서늘한 곳에 보관'}
    usage_map = {'냉동': '소비자 판매용 / 해동 후 섭취', '냉장': '소비자 판매용 / 냉장 보관 후 섭취', '실온': '소비자 판매용 / 개봉 후 섭취'}

    # 비교제품용 주성분
    top3 = sorted_ings[:3]
    top3_names = []
    for ing in top3:
        info = find_ingredient(ing['name'], db)
        top3_names.append(info.get('원재료명', ing['name']) if info else ing['name'])
    main_ingredients = ", ".join(top3_names)

    compare_product = options.get('compare_product', '통밀식빵(냉동)')
    compare_company = options.get('compare_company', '더브레드블루')

    today = datetime.now()

    # 임시 디렉토리
    tmp_dir = tempfile.mkdtemp()
    output_files = {}

    try:
        hwp = win32com.client.gencache.EnsureDispatch('HWPFrame.HwpObject')
        hwp.XHwpWindows.Item(0).Visible = False
        hwp.RegisterModule('FilePathCheckDLL', 'FilePathCheckerModule')

        # ─── 1. 제조방법설명서 ───
        shutil.copy2(HWP_TEMPLATES['doc1'], os.path.join(tmp_dir, 'doc1.hwp'))
        hwp.Open(os.path.join(tmp_dir, 'doc1.hwp'), 'HWP', 'forceopen:true')

        hwp_replace(hwp, '호라산밀식빵', product_name)
        # 원재료명 배합비율 — 기존 긴 텍스트를 통째로 교체
        hwp_replace(hwp, '밀가루 50.5%, 정제수 33.19%, 가공두유 4.81%, 미강유(현미유) 2.81%, 곡류가공품1 2.16%, 혼합제제 1.7%, 전분가공품 1.44%, 식품첨가물(효모) 1.44%, 곡류가공품2 0.82%, 정제소금 0.62%, 당류가공품 0.51%',
                    ingredient_str)
        # 제조방법
        hwp_replace(hwp, '호라산밀가루,정제수,두유,현미유,찹쌀가루,스테비아,타피오카전분,효모,엔지마띠코,정제소금,몰트엑기스를 계량후 믹싱을 한다', method_text.split('\r\n')[0])
        hwp_replace(hwp, '(1)을 1차발효 40~50분을 한다', f'(1)을 1차발효 {ferment1}분을 한다')
        hwp_replace(hwp, '(2)를 320g 분할후 둥글리기를 한다', f'(2)를 {div_wt}g 분할후 둥글리기를 한다')
        hwp_replace(hwp, '(3)을 식빵팬에 넣고 2차발효 50~60분을 한다', f'(3)을 팬에 넣고 2차발효 {ferment2}분을 한다')
        hwp_replace(hwp, '(4)를 180℃로 예열된 오븐에 넣고 45~55분간 가열을 한다', f'(4)를 {bake_temp}℃로 예열된 오븐에 넣고 {bake_time}분간 가열을 한다')
        hwp_replace(hwp, '(5)를 탈팬후 식혀 포장 냉동을 한다', f'(5)를 탈팬후 식혀 {storage_suffix.get(storage_method, storage_suffix["냉동"])}')
        hwp_replace(hwp, '소비자 판매용 / 해동 후 섭취', usage_map.get(storage_method, usage_map['냉동']))
        hwp_replace(hwp, '-18℃이하 냉동보관', storage_map.get(storage_method, storage_map['냉동']))
        hwp_replace(hwp, 'PP, PS, PE, PET', packaging_material)
        hwp_replace(hwp, '밀봉 / 10g~5kg', f"밀봉 / {options.get('packaging_unit', '10g~5kg')}")
        hwp_replace(hwp, '제조일로부터 9개월까지', f'제조일로부터 {shelf_life}까지')

        save1 = os.path.join(tmp_dir, f'1. 제조방법설명서 - {product_name}.hwp')
        hwp.HAction.GetDefault('FileSaveAs_S', hwp.HParameterSet.HFileOpenSave.HSet)
        hwp.HParameterSet.HFileOpenSave.filename = save1
        hwp.HParameterSet.HFileOpenSave.Format = 'HWP'
        hwp.HAction.Execute('FileSaveAs_S', hwp.HParameterSet.HFileOpenSave.HSet)
        hwp.HAction.Run('FileClose')

        # ─── 2. 원료성분 및 배합비율 ───
        shutil.copy2(HWP_TEMPLATES['doc2'], os.path.join(tmp_dir, 'doc2.hwp'))
        hwp.Open(os.path.join(tmp_dir, 'doc2.hwp'), 'HWP', 'forceopen:true')

        # 기존 원재료 데이터를 새 데이터로 교체 (찾아바꾸기)
        # 호라산밀식빵의 기존 원재료들
        old_ings = [
            ('호라산밀', '50.50', '유기농호라산밀100%', '호주'),
            ('정제수', '33.19', '', ''),
            ('두유', '4.81', '원액두유99.9%[대두고형분10%이상,외국산\r\n(미국,캐니다,러시아등)],식염', ''),
            ('현미유', '2.81', '미강유100%', '태국'),
            ('찹쌀가루', '2.16', '찹쌀100%', '국내산'),
            ('스테비아', '1.70', '에리스리톨97.9%,효소처리스테비아2.1%', ''),
            ('타피오카전분', '1.44', '변성전분99.9%(태국산),감자전분0.1%', ''),
            ('효모', '1.44', '효모(사카로마이세스 세레비제)98.5%, 소르비탄지방산에스테르1.5%', '프랑스'),
            ('엔지마띠코', '0.82', '건조소맥분,효소(자일라나아제,알파아밀라아제\r\n(비세균성),리파아제,글루코아밀라아제)', ''),
            ('정제소금', '0.62', '해수100%', ''),
            ('몰트엑기스', '0.51', '물엿,정제수,맥아엑기스8%[(덴마크산)\r\n고형분67%이상]', ''),
        ]

        # 기존 원재료명들을 새 데이터로 교체
        for i, old in enumerate(old_ings):
            if i < len(sorted_ings):
                ing = sorted_ings[i]
                info = find_ingredient(ing['name'], db)
                new_name = info.get('원재료명', ing['name']) if info else ing['name']
                new_pct = f"{ing['percent']:.2f}"
                new_comp = info.get('성분', '') if info else ''
                new_origin = info.get('원산지', '') if info else ''

                hwp_replace(hwp, old[0], new_name)
                hwp_replace(hwp, old[1], new_pct)
            else:
                # 원재료가 더 적으면 빈칸으로
                hwp_replace(hwp, old[0], '')
                hwp_replace(hwp, old[1], '')

        # 성분과 비고는 복잡하므로, 전체 텍스트를 한번에 교체
        for old in old_ings:
            if old[2]:
                clean_old = old[2].replace('\r\n', '')
                hwp_replace(hwp, clean_old, '')

        # 새 성분/비고 데이터는 이름 옆에 이미 매칭되어 있으므로
        # 각 원재료 셀 뒤의 성분 셀에 새 데이터 입력
        for i, ing in enumerate(sorted_ings):
            info = find_ingredient(ing['name'], db)
            if info:
                new_name = info.get('원재료명', ing['name'])
                new_comp = info.get('성분', '')
                new_origin = info.get('원산지', '')
                # 원재료명이 이미 교체되었으므로, 해당 이름 뒤에 있는 빈 셀에 성분 입력
                # 찾아바꾸기로는 한계가 있어서, 전체 텍스트 방식으로

        save2 = os.path.join(tmp_dir, f'2. 원료성분 및 배합비율({product_name}).hwp')
        hwp.HAction.GetDefault('FileSaveAs_S', hwp.HParameterSet.HFileOpenSave.HSet)
        hwp.HParameterSet.HFileOpenSave.filename = save2
        hwp.HParameterSet.HFileOpenSave.Format = 'HWP'
        hwp.HAction.Execute('FileSaveAs_S', hwp.HParameterSet.HFileOpenSave.HSet)
        hwp.HAction.Run('FileClose')

        # ─── 3. 소비기한설정사유서 ───
        shutil.copy2(HWP_TEMPLATES['doc3'], os.path.join(tmp_dir, 'doc3.hwp'))
        hwp.Open(os.path.join(tmp_dir, 'doc3.hwp'), 'HWP', 'forceopen:true')

        hwp_replace(hwp, '호라산밀식빵', product_name)
        hwp_replace(hwp, '제조일로부터 9개월까지', f'제조일로부터 {shelf_life}까지')
        hwp_replace(hwp, '통밀식빵(냉동)', compare_product)

        # 보존방법 체크 변경
        if storage_method == '냉동':
            pass  # 기본값
        elif storage_method == '냉장':
            hwp_replace(hwp, '냉동( O )', '냉동(   )')
            hwp_replace(hwp, '냉장(   )', '냉장( O )')
        elif storage_method == '실온':
            hwp_replace(hwp, '냉동( O )', '냉동(   )')
            hwp_replace(hwp, '실온(   )', '실온( O )')

        # 주성분
        hwp_replace(hwp, '밀가루, 정제수, 가공두유', main_ingredients)
        hwp_replace(hwp, 'PP, PS, PE, PET', packaging_material)

        # 날짜
        hwp_replace(hwp, '2025년    04월   25일', f'{today.year}년    {today.month:02d}월    {today.day:02d}일')

        save3 = os.path.join(tmp_dir, f'3. 소비기한설정사유서({product_name}).hwp')
        hwp.HAction.GetDefault('FileSaveAs_S', hwp.HParameterSet.HFileOpenSave.HSet)
        hwp.HParameterSet.HFileOpenSave.filename = save3
        hwp.HParameterSet.HFileOpenSave.Format = 'HWP'
        hwp.HAction.Execute('FileSaveAs_S', hwp.HParameterSet.HFileOpenSave.HSet)
        hwp.HAction.Run('FileClose')

        hwp.Quit()

        # 파일 읽기
        for key, path in [('hwp1', save1), ('hwp2', save2), ('hwp3', save3)]:
            with open(path, 'rb') as f:
                output_files[key] = f.read()

    except Exception as e:
        try:
            hwp.Quit()
        except:
            pass
        raise e
    finally:
        pythoncom.CoUninitialize()
        # 임시 파일 정리
        try:
            shutil.rmtree(tmp_dir, ignore_errors=True)
        except:
            pass

    return output_files

# ─── Streamlit UI ───
st.title("📋 품목제조보고서 자동생성")
st.caption("레시피(원가) 엑셀 파일을 업로드하면 품목제조보고서 3종을 자동 생성합니다.")

db = load_ingredient_db()

# 파일 업로드
uploaded_file = st.file_uploader(
    "레시피(원가) 엑셀 파일을 드래그앤드롭하세요",
    type=['xlsx', 'xls'],
    help="원재료명, 1배합, 백분율 컬럼이 있는 엑셀 파일"
)

if uploaded_file:
    recipes = parse_recipe_excel(uploaded_file)

    if not recipes:
        st.error("레시피 데이터를 찾을 수 없습니다. 엑셀 파일 형식을 확인해주세요.")
        st.stop()

    # 시트 선택 (여러 시트가 있을 경우)
    if len(recipes) > 1:
        sheet_names = [f"{r['product_name']} ({r['sheet_name']})" for r in recipes]
        selected = st.selectbox("레시피 시트 선택", sheet_names)
        recipe = recipes[sheet_names.index(selected)]
    else:
        recipe = recipes[0]

    st.subheader(f"📦 {recipe['product_name']}")

    # 별칭 매핑 로드
    if 'alias_map' not in st.session_state:
        st.session_state.alias_map = load_alias_map()

    # 원재료 미리보기
    with st.expander("원재료 배합비율 확인", expanded=True):
        cols_header = st.columns([3, 2, 4, 2])
        cols_header[0].markdown("**원재료명**")
        cols_header[1].markdown("**배합비율(%)**")
        cols_header[2].markdown("**성분**")
        cols_header[3].markdown("**DB매칭**")

        unmatched = []
        for ing in sorted(recipe['ingredients'], key=lambda x: x['percent'], reverse=True):
            info = find_ingredient(ing['name'], db)
            cols = st.columns([3, 2, 4, 2])
            cols[0].write(ing['name'])
            cols[1].write(f"{ing['percent']:.2f}%")
            if info:
                cols[2].write(info.get('성분', '-'))
                cols[3].write("✅")
            else:
                cols[2].write("⚠️ DB 미등록")
                cols[3].write("❌")
                unmatched.append(ing['name'])

    # 미매칭 원재료 매칭 UI
    if unmatched:
        st.divider()
        st.subheader("🔍 원재료 매칭")
        st.info("DB에 없는 원재료가 있습니다. 아래에서 올바른 매칭을 선택해주세요.")

        alias_changed = False
        for um_name in unmatched:
            suggestions = suggest_matches(um_name, db, top_n=5)
            top_match = suggestions[0] if suggestions else None

            st.markdown(f"---")
            st.markdown(f"**'{um_name}'** 에 대한 추천:")

            # 자동 추천 표시
            if top_match and top_match[1] > 0.3:
                good_suggestions = [s for s in suggestions if s[1] > 0.15]
                for s in good_suggestions:
                    st.caption(f"  ✨ **{s[0]}** (유사도 {s[1]:.0%}) — {s[2][:50]}")

            # 검색 + 선택 UI (타이핑하면 필터링)
            search_text = st.text_input(
                f"🔍 원재료 검색",
                value="",
                key=f"search_{um_name}",
                placeholder=f"검색어 입력 (예: 밀가루, 버터, 올리브...)"
            )

            if search_text:
                # 검색어로 DB 필터링
                search_lower = search_text.strip()
                filtered = {}
                for key, val in db.items():
                    searchable = f"{key} {val.get('원재료명','')} {val.get('성분','')} {val.get('식품유형','')}"
                    if search_lower in searchable:
                        filtered[key] = val

                if filtered:
                    # 검색 결과를 selectbox로 표시
                    result_options = [f"{k} — {v.get('성분','')[:40]}" for k, v in filtered.items()]
                    selected_result = st.selectbox(
                        f"검색 결과 ({len(filtered)}건)",
                        result_options,
                        key=f"result_{um_name}"
                    )
                    if selected_result:
                        matched_key = selected_result.split(" — ")[0].strip()
                        st.caption(f"성분: {db[matched_key].get('성분', '-')} | 원산지: {db[matched_key].get('원산지', '-')}")
                        if st.button(f"✅ '{um_name}' → '{matched_key}' 매칭 확인", key=f"confirm_{um_name}"):
                            st.session_state.alias_map[um_name] = matched_key
                            save_alias_map(st.session_state.alias_map)
                            st.success(f"'{um_name}' → '{matched_key}' 매칭 저장!")
                            alias_changed = True
                else:
                    st.caption("검색 결과 없음")
            elif top_match and top_match[1] > 0.3:
                # 추천에서 바로 선택
                best_key = suggestions[0][0]
                if st.button(f"✅ '{um_name}' → '{best_key}' 매칭 확인", key=f"confirm_top_{um_name}"):
                    st.session_state.alias_map[um_name] = best_key
                    save_alias_map(st.session_state.alias_map)
                    st.success(f"'{um_name}' → '{best_key}' 매칭 저장!")
                    alias_changed = True

            # 직접 입력
            with st.expander(f"'{um_name}' 성분 직접 입력", expanded=False):
                manual_comp = st.text_input("성분", key=f"manual_{um_name}", placeholder="예: 올리브 100%")
                manual_origin = st.text_input("원산지", key=f"origin_{um_name}", placeholder="예: 미국")
                if manual_comp and st.button(f"DB에 추가", key=f"save_manual_{um_name}"):
                    db[um_name] = {
                        "식품유형": "",
                        "원재료명": um_name,
                        "표시명": um_name,
                        "성분": manual_comp,
                        "원산지": manual_origin,
                        "알레르기": ""
                    }
                    db_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "ingredient_db.json")
                    with open(db_path, "w", encoding="utf-8") as f:
                        json.dump(db, f, ensure_ascii=False, indent=2)
                    st.success(f"'{um_name}' DB에 추가 완료!")
                    load_ingredient_db.clear()
                    alias_changed = True

        if alias_changed:
            st.warning("매칭 정보가 변경되었습니다. 새로고침하면 반영됩니다.")
            st.button("🔄 새로고침", on_click=lambda: st.rerun())

    # 저장된 별칭 매핑 표시
    if st.session_state.alias_map:
        with st.expander(f"저장된 별칭 매핑 ({len(st.session_state.alias_map)}건)"):
            for alias, target in st.session_state.alias_map.items():
                c1, c2, c3 = st.columns([3, 3, 1])
                c1.write(f"'{alias}'")
                c2.write(f"→ '{target}'")
                if c3.button("삭제", key=f"del_alias_{alias}"):
                    del st.session_state.alias_map[alias]
                    save_alias_map(st.session_state.alias_map)
                    st.rerun()

    st.divider()

    # 옵션 설정
    st.subheader("⚙️ 제조 옵션 설정")

    col1, col2 = st.columns(2)

    with col1:
        storage = st.selectbox("보관방법", ['냉동', '냉장', '실온'], index=0)
        shelf_life = st.selectbox("소비기한", [
            '3개월', '6개월', '9개월', '12개월', '15개월', '18개월', '24개월'
        ], index=2)
        packaging_material = st.text_input("포장재질", value="PP, PS, PE, PET")
        packaging_unit = st.text_input("포장단위", value="10g~5kg")

    with col2:
        bake_temp = st.text_input("굽기 온도(℃)", value="180")
        bake_time = st.text_input("굽기 시간(분)", value="30~40")
        ferment1_time = st.text_input("1차발효 시간(분)", value="40~50")
        ferment2_time = st.text_input("2차발효 시간(분)", value="50~60")
        divide_weight = st.text_input("분할 중량(g)", value="320")

    st.divider()

    # 비교제품 설정 (소비기한설정사유서용) — 자사 기존 제품에서 검색/선택
    st.subheader("📊 비교제품 설정 (소비기한설정사유서)")

    # 품목제조보고 폴더에서 기존 제품 목록 로드
    @st.cache_data
    def load_existing_products():
        # 먼저 JSON 파일에서 로드 (클라우드 호환)
        json_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'existing_products.json')
        if os.path.exists(json_path):
            with open(json_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        # 로컬: 폴더에서 직접 스캔
        base = r'c:/Users/Moon/OneDrive/03. 양지희_품질/2. 품목제조보고/1. 빵류'
        products = set()
        if os.path.exists(base):
            for year_folder in os.listdir(base):
                year_path = os.path.join(base, year_folder)
                if not os.path.isdir(year_path):
                    continue
                for pf in os.listdir(year_path):
                    if os.path.isdir(os.path.join(year_path, pf)):
                        clean = pf.strip()
                        if clean and not clean.startswith('★') and not clean.startswith('260') and not clean.startswith('New'):
                            products.add(clean)
        return sorted(products)

    existing_products = load_existing_products()

    compare_search = st.text_input(
        "🔍 비교제품 검색 (자사 기존 제품)",
        placeholder="검색어 입력 (예: 통밀, 식빵, 베이글...)",
        key="compare_search"
    )

    if compare_search:
        filtered_products = [p for p in existing_products if compare_search in p]
        if filtered_products:
            compare_product = st.selectbox(
                f"비교제품 선택 ({len(filtered_products)}건)",
                filtered_products,
                key="compare_select"
            )
        else:
            st.caption("검색 결과 없음 — 직접 입력하세요")
            compare_product = st.text_input("비교제품명 직접 입력", value="통밀식빵(냉동)", key="compare_manual")
    else:
        compare_product = st.selectbox(
            f"비교제품 선택 (전체 {len(existing_products)}건)",
            existing_products,
            index=existing_products.index("순수 통밀식빵") if "순수 통밀식빵" in existing_products else 0,
            key="compare_all"
        )

    compare_company = st.text_input("비교제품 제조사", value="더브레드블루")

    st.divider()

    options = {
        'storage': storage,
        'shelf_life': shelf_life,
        'packaging_material': packaging_material,
        'packaging_unit': packaging_unit,
        'bake_temp': bake_temp,
        'bake_time': bake_time,
        'ferment1_time': ferment1_time,
        'ferment2_time': ferment2_time,
        'divide_weight': divide_weight,
        'compare_product': compare_product,
        'compare_company': compare_company,
    }

    # 생성 버튼 — 결과를 session_state에 저장해서 다운로드 시 사라지지 않게
    if st.button("📝 품목제조보고서 생성", type="primary", use_container_width=True):
        with st.spinner("문서 생성 중... (HWP 생성에 10~20초 소요)"):
            product_name = recipe['product_name']

            # Word 문서 생성
            doc1 = generate_doc1_manufacturing(recipe, options, db)
            doc2 = generate_doc2_ingredients(recipe, options, db)
            doc3 = generate_doc3_shelf_life(recipe, options, db)

            # PDF 생성
            pdf1 = generate_pdf('manufacturing', recipe, options, db)
            pdf2 = generate_pdf('ingredients', recipe, options, db)
            pdf3 = generate_pdf('shelf_life', recipe, options, db)

            # Word → bytes
            buf1 = io.BytesIO(); doc1.save(buf1); buf1.seek(0)
            buf2 = io.BytesIO(); doc2.save(buf2); buf2.seek(0)
            buf3 = io.BytesIO(); doc3.save(buf3); buf3.seek(0)

            # HWP 생성 (로컬 Windows + 한글 설치 시에만)
            hwp_files = None
            if os.name == 'nt':
                try:
                    hwp_files = generate_hwp_files(recipe, options, db)
                except Exception as e:
                    hwp_files = None
                    st.warning(f"HWP 생성 실패 (한글 프로그램 필요): {e}")

            # session_state에 저장
            st.session_state.generated_files = {
                'product_name': product_name,
                'doc1': buf1.getvalue(),
                'doc2': buf2.getvalue(),
                'doc3': buf3.getvalue(),
                'pdf1': pdf1.getvalue(),
                'pdf2': pdf2.getvalue(),
                'pdf3': pdf3.getvalue(),
            }
            if hwp_files:
                st.session_state.generated_files.update(hwp_files)

        st.rerun()

    # 생성된 파일이 있으면 다운로드 UI 표시 (rerun 후에도 유지됨)
    if 'generated_files' in st.session_state:
        files = st.session_state.generated_files
        product_name = files['product_name']

        st.success("✅ 문서 생성 완료!")

        st.subheader("📥 다운로드")

        # Word 다운로드
        st.markdown("**Word 파일 (.docx)**")
        wcol1, wcol2, wcol3 = st.columns(3)

        wcol1.download_button(
            "1. 제조방법설명서",
            files['doc1'],
            f"1. 제조방법설명서 - {product_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        wcol2.download_button(
            "2. 원료성분 및 배합비율",
            files['doc2'],
            f"2. 원료성분 및 배합비율({product_name}).docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        wcol3.download_button(
            "3. 소비기한설정사유서",
            files['doc3'],
            f"3. 소비기한설정사유서({product_name}).docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        st.markdown("---")

        # PDF 다운로드
        st.markdown("**PDF 파일 (제출용)**")
        pcol1, pcol2, pcol3 = st.columns(3)

        pcol1.download_button(
            "1. 제조방법설명서 (PDF)",
            files['pdf1'],
            f"1. 제조방법설명서 - {product_name}.pdf",
            mime="application/pdf"
        )

        pcol2.download_button(
            "2. 원료성분 및 배합비율 (PDF)",
            files['pdf2'],
            f"2. 원료성분 및 배합비율({product_name}).pdf",
            mime="application/pdf"
        )

        pcol3.download_button(
            "3. 소비기한설정사유서 (PDF)",
            files['pdf3'],
            f"3. 소비기한설정사유서({product_name}).pdf",
            mime="application/pdf"
        )

        # HWP 다운로드
        if 'hwp1' in files:
            st.markdown("---")
            st.markdown("**한글 파일 (.hwp) — 기존 양식 그대로**")
            hcol1, hcol2, hcol3 = st.columns(3)

            hcol1.download_button(
                "1. 제조방법설명서 (HWP)",
                files['hwp1'],
                f"1. 제조방법설명서 - {product_name}.hwp",
                mime="application/x-hwp"
            )

            hcol2.download_button(
                "2. 원료성분 및 배합비율 (HWP)",
                files['hwp2'],
                f"2. 원료성분 및 배합비율({product_name}).hwp",
                mime="application/x-hwp"
            )

            hcol3.download_button(
                "3. 소비기한설정사유서 (HWP)",
                files['hwp3'],
                f"3. 소비기한설정사유서({product_name}).hwp",
                mime="application/x-hwp"
            )

        st.divider()
        st.info("💡 영양성분 정보가 필요하면 [영양성분 자동계산기](https://thebreadblue-nutrition-zh7cydsadtwdmymbyawsbz.streamlit.app/)를 이용하세요.")

else:
    st.info("👆 위에 엑셀 파일을 업로드하면 시작됩니다.")

    with st.expander("사용 방법"):
        st.markdown("""
        1. **레시피(원가) 엑셀 파일**을 드래그앤드롭합니다
           - 원재료명, 1배합(g), 백분율(%) 컬럼이 있어야 합니다
        2. 원재료 매칭 결과를 확인합니다
        3. **보관방법, 소비기한, 제조조건** 등을 설정합니다
        4. **품목제조보고서 생성** 버튼을 클릭합니다
        5. Word(편집용) + PDF(제출용) 6개 파일을 다운로드합니다

        **생성되는 문서:**
        - 1. 제조방법설명서
        - 2. 원료성분 및 배합비율
        - 3. 소비기한설정사유서
        """)

    with st.expander(f"등록된 원재료 DB ({len(db)}종)"):
        for name, info in sorted(db.items()):
            st.write(f"- **{name}**: {info.get('성분', '-')} ({info.get('원산지', '')})")
